import re
import io
import asyncio
from typing import List, Dict, Optional, Tuple
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel, Field
import httpx
from rapidfuzz import fuzz

# PDF
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch

# DOCX
from docx import Document
from docx.shared import Pt

# -----------------------------------------------------------------------------
# APP CONFIG
# -----------------------------------------------------------------------------
app = FastAPI(
    title="AI Detector PRO",
    version="3.2.0",
    openapi_url="/openapi.json",
    docs_url="/docs",
    redoc_url="/redoc",
)

@app.get("/health")
def health():
    return {"ok": True}

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
    allow_credentials=True,
    allow_methods=["*"],
)

# -----------------------------------------------------------------------------
# UTILS
# -----------------------------------------------------------------------------
_SENTENCE_RE = re.compile(r'(?us)([^.!?]+[.!?])')
_TAG_RE = re.compile(r"<[^>]+>")
_WS_RE = re.compile(r"\s+")

HTTP_LIMITS = httpx.Limits(max_connections=5, max_keepalive_connections=5)
HTTP_TIMEOUT = httpx.Timeout(connect=5.0, read=8.0, write=8.0, pool=5.0)

def clean_text(text: str) -> str:
    text = _TAG_RE.sub(" ", text)
    return _WS_RE.sub(" ", text).strip()

def split_sentences(text: str) -> List[str]:
    parts = _SENTENCE_RE.findall(text) or [text]
    return [p.strip() for p in parts if p.strip()]

def stitch(sents: List[str]) -> str:
    return " ".join(s for s in sents if s.strip())

def ai_score(text: str) -> int:
    sents = split_sentences(text)
    long_count = sum(len(s.split()) >= 18 for s in sents)
    return min(100, 15 + long_count * 12)

async def fetch(url: str, client: httpx.AsyncClient) -> str:
    try:
        r = await client.get(url)
        if r.status_code == 200:
            return clean_text(r.text)
    except:
        pass
    return ""

async def build_corpus(texts: List[str], urls: List[str]) -> List[Tuple[str, str]]:
    corpus = [(f"text:{i}", clean_text(t)) for i, t in enumerate(texts or []) if (t or "").strip()]
    if urls:
        # cap to keep latency under control
        urls = urls[:5]
        async with httpx.AsyncClient(
            follow_redirects=True, limits=HTTP_LIMITS, timeout=HTTP_TIMEOUT
        ) as client:
            res = await asyncio.gather(*[fetch(u, client) for u in urls])
        for u, body in zip(urls, res):
            if body:
                corpus.append((u, body))
    return corpus

def similarity(s: str, body: str) -> Tuple[int, str]:
    best, snip = 0, ""
    for cs in split_sentences(body):
        sc = max(fuzz.token_set_ratio(s, cs), fuzz.partial_ratio(s, cs))
        if sc > best:
            best, snip = sc, cs
    return best, snip

def reason_ai(s: str) -> str:
    long = len(s.split()) >= 22
    passive = any(x in s.lower() for x in ["is being", "was", "were", "has been", "have been"])
    tokens = [w.strip(",.!?;:").lower() for w in s.split()]
    lowvar = len(set(tokens)) / max(1, len(tokens)) < 0.65
    r = []
    if long: r.append("very long")
    if passive: r.append("passive tone")
    if lowvar: r.append("low lexical variety")
    return ", ".join(r) if r else "generic phrasing"

# ---------------- Better rule stack for rewrites ----------------
def rewrite_rules(sentence: str, style: str) -> str:
    s = re.sub(r"\s+", " ", sentence.strip())

    if style == "concise":
        s = re.sub(r"\b(in order to|due to the fact that|with respect to|it is important to note that)\b", "to", s, flags=re.I)
        s = re.sub(r"\b(utilize)\b", "use", s, flags=re.I)
        # trim long chains
        words = s.split()
        if len(words) > 28:
            s = " ".join(words[:24]) + "."
        return s
    if style == "formal":
        s = re.sub(r"\b(can't|won't|don't|isn't|aren't|doesn't)\b",
                   lambda m: {"can't":"cannot","won't":"will not","don't":"do not",
                              "isn't":"is not","aren't":"are not","doesn't":"does not"}[m.group(0)], s, flags=re.I)
        s = re.sub(r"\b(kids)\b", "children", s, flags=re.I)
        return s
    if style == "simple":
        s = re.sub(r",\s*", ". ", s)  # split into shorter thoughts
        s = re.sub(r"\b(utilize)\b", "use", s, flags=re.I)
        return s
    # humanize (default)
    s = re.sub(r"\b(is|was|were|been|being|are) ([a-z]+)ed\b", r"\2", s)  # reduce passive
    if s.endswith("."): s = s[:-1]
    return s + " with a specific example."

# -----------------------------------------------------------------------------
# MODELS
# -----------------------------------------------------------------------------
class AnalyzeIn(BaseModel):
    text: str
    source_texts: Optional[List[str]] = None
    source_urls: Optional[List[str]] = None
    plag_threshold: int = Field(default=85, ge=50, le=100)

class SentOut(BaseModel):
    index: int
    text: str
    label: str                # ai | plag | human
    confidence: str           # Low | Medium | High
    plag_score: Optional[int] = None
    source: Optional[str] = None
    snippet: Optional[str] = None
    citation_index: Optional[int] = None
    reason: Optional[str] = None
    fix: Optional[str] = None

class AnalyzeOut(BaseModel):
    summary: dict
    counts: dict
    sentences: List[SentOut]
    citations: List[str]         # ordered list of sources -> [1], [2], ...
    markdown_report: str

# -----------------------------------------------------------------------------
# CORE ANALYSIS (with citations)
# -----------------------------------------------------------------------------
async def core(req: AnalyzeIn) -> AnalyzeOut:
    text = req.text.strip()
    sents = split_sentences(text)
    ai_over = ai_score(text)
    corpus = await build_corpus(req.source_texts or [], req.source_urls or [])

    # Citation map for URLs/texts
    cite_map: Dict[str, int] = {}
    ordered_citations: List[str] = []

    out: List[SentOut] = []
    ai_n = pl_n = 0

    seen = set()
    for i, s in enumerate(sents):
        isdup = s in seen
        seen.add(s)

        best, src, snip = 0, None, None
        for u, b in corpus:
            sc, mp = similarity(s, b)
            if sc > best:
                best, src, snip = sc, u, mp

        label = "human"
        if best >= req.plag_threshold:
            label = "plag"
        elif len(s.split()) >= 18 or ai_over >= 70:
            label = "ai"

        if label == "ai": ai_n += 1
        if label == "plag": pl_n += 1

        # confidence
        conf = "High" if (label == "plag" and best >= 92) or (label == "ai" and ai_over >= 80) else ("Medium" if label != "human" else "Low")

        # citation index
        cidx = None
        if label == "plag" and src:
            if src not in cite_map:
                cite_map[src] = len(ordered_citations) + 1
                ordered_citations.append(src)
            cidx = cite_map[src]

        out.append(SentOut(
            index=i, text=s, label=label, confidence=conf,
            plag_score=best if best else None, source=src, snippet=snip,
            citation_index=cidx,
            reason=reason_ai(s) if label == "ai" else ("high similarity" if label == "plag" else None),
            fix=("Add personal detail & varied verbs." if label == "ai" else ("Paraphrase in your words & cite." if label == "plag" else None))
        ))

    total = len(out)
    plag_pct = min(100, int(100 * pl_n / max(1, total)))
    summary = {
        "ai_score": ai_over,
        "plagiarism_score": plag_pct,
        "human_score": max(0, 100 - max(ai_over, plag_pct)),
    }
    counts = {"total": total, "ai": ai_n, "plag": pl_n, "human": total - ai_n - pl_n}

    # Markdown report with footnotes
    md = ["## AI Text Detection Report",
          f"- **AI-likeness:** {summary['ai_score']}%",
          f"- **Plagiarism risk:** {summary['plagiarism_score']}%",
          f"- **Human-likeness:** {summary['human_score']}%",
          "⚠️ Indicators only. Verify sources & context."]
    if ai_n:
        md.append("\n### AI-flagged sentences")
        for s in out:
            if s.label == "ai":
                md.append(f"**{s.index+1}.** {s.text}\n  - Reason: {s.reason}\n  - Confidence: {s.confidence}\n  - Fix: {s.fix}")
    else:
        md.append("\n### AI-flagged sentences\nNone detected ✅")

    if pl_n:
        md.append("\n### Plagiarism-risk sentences")
        for s in out:
            if s.label == "plag":
                tag = f" [^{s.citation_index}]" if s.citation_index else ""
                sn = f'  - Match: “{s.snippet}”' if s.snippet else ""
                md.append(f"**{s.index+1}.** {s.text}{tag}\n  - Similarity: {s.plag_score}% | Confidence: {s.confidence}\n{sn}\n  - Fix: {s.fix}")
        if ordered_citations:
            md.append("\n### Sources")
            for i, u in enumerate(ordered_citations, 1):
                md.append(f"[^{i}]: {u}")
    else:
        md.append("\n### Plagiarism-risk sentences\nNone detected ✅")

    return AnalyzeOut(summary=summary, counts=counts, sentences=out, citations=ordered_citations, markdown_report="\n".join(md))

# -----------------------------------------------------------------------------
# API ENDPOINTS
# -----------------------------------------------------------------------------
@app.post("/analyze", response_model=AnalyzeOut)
async def analyze(req: AnalyzeIn):
    return JSONResponse((await core(req)).model_dump())

class RewriteIn(BaseModel):
    sentence: str
    style: str = Field(default="humanize", description="humanize | concise | formal | simple")

@app.post("/rewrite")
def rewrite(req: RewriteIn):
    style = req.style.lower()
    if style not in {"humanize", "concise", "formal", "simple"}:
        style = "humanize"
    out = rewrite_rules(req.sentence, style)
    tip = {
        "humanize": "Add concrete detail, varied verbs, and first-person perspective where appropriate.",
        "concise": "Drop filler phrases, shorten clauses, prefer strong verbs over adverbs.",
        "formal": "Use full forms, neutral tone, precise vocabulary.",
        "simple": "Use short sentences and common words; avoid jargon."
    }[style]
    return {"original": req.sentence, "style": style, "rewritten": out, "tips": tip}

class RewriteBatchIn(BaseModel):
    text: str
    style_map: Optional[Dict[str, str]] = Field(default=None, description="per-label style {ai,plag,human}")
    default_style: str = "humanize"
    plag_threshold: int = 85
    source_texts: Optional[List[str]] = None
    source_urls: Optional[List[str]] = None

@app.post("/rewrite-batch")
async def rewrite_batch(req: RewriteBatchIn):
    analysis = await core(AnalyzeIn(
        text=req.text, source_texts=req.source_texts, source_urls=req.source_urls, plag_threshold=req.plag_threshold
    ))
    styles = {"ai": req.default_style, "plag": req.default_style, "human": req.default_style}
    if req.style_map:
        styles.update({k: v for k, v in req.style_map.items() if k in styles})

    originals, rewrites = [], []
    flagged_only = []
    all_new = []
    for s in analysis.sentences:
        st = styles.get(s.label, req.default_style)
        new = rewrite_rules(s.text, st)
        originals.append({"index": s.index, "label": s.label, "text": s.text})
        rewrites.append({"index": s.index, "label": s.label, "style": st, "rewritten": new})
        flagged_only.append(new if s.label in {"ai", "plag"} else s.text)
        all_new.append(new)

    return {
        "replaced_count": sum(1 for s in analysis.sentences if s.label in {"ai", "plag"}),
        "originals": originals,
        "rewrites": rewrites,
        "rewritten_text_flagged_only": stitch([x for x in flagged_only]),
        "rewritten_text_all": stitch([x for x in all_new]),
        "citations": analysis.citations,
        "summary": analysis.summary
    }

class PdfIn(BaseModel):
    text: str
    filename: str = "detector_report.pdf"
    source_texts: Optional[List[str]] = None
    source_urls: Optional[List[str]] = None
    plag_threshold: int = 85

@app.post("/report.pdf")
async def report_pdf(req: PdfIn):
    analysis = await core(AnalyzeIn(
        text=req.text, source_texts=req.source_texts, source_urls=req.source_urls, plag_threshold=req.plag_threshold
    ))

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    W, H = letter
    left = 1*inch
    width = W - 2*inch
    y = H - 1*inch

    def line(txt, bold=False, size=10, pad=14):
        nonlocal y
        c.setFont("Helvetica-Bold" if bold else "Helvetica", size)
        c.drawString(left, y, txt)
        y -= pad
        if y < 1*inch:
            c.showPage(); y = H - 1*inch

    def wrap(txt, size=10, pad=14):
        nonlocal y
        c.setFont("Helvetica", size)
        words = txt.split(" ")
        row = ""
        for w in words:
            probe = (row + " " + w).strip()
            if c.stringWidth(probe, "Helvetica", size) <= width:
                row = probe
            else:
                c.drawString(left, y, row); y -= pad; row = w
                if y < 1*inch: c.showPage(); y = H - 1*inch
        if row:
            c.drawString(left, y, row); y -= pad
        if y < 1*inch: c.showPage(); y = H - 1*inch

    line("AI Text Detection Report", bold=True, size=16, pad=20)
    wrap("This report is informational. Use human judgment and cite sources where appropriate.", size=9, pad=12)

    line("Overall Results", bold=True, size=12, pad=16)
    wrap(f"AI-likeness: {analysis.summary['ai_score']}%")
    wrap(f"Plagiarism risk: {analysis.summary['plagiarism_score']}%")
    wrap(f"Human-likeness: {analysis.summary['human_score']}%")

    line("AI-flagged sentences", bold=True, size=12, pad=16)
    ai_items = [s for s in analysis.sentences if s.label == "ai"]
    if not ai_items:
        wrap("None detected ✅")
    else:
        for s in ai_items:
            wrap(f"{s.index+1}. {s.text}")
            wrap(f"Reason: {s.reason} | Confidence: {s.confidence}")
            wrap(f"Fix: {s.fix}", size=9, pad=12)

    line("Plagiarism-risk sentences", bold=True, size=12, pad=16)
    pl_items = [s for s in analysis.sentences if s.label == "plag"]
    if not pl_items:
        wrap("None detected ✅")
    else:
        for s in pl_items:
            tag = f" [{s.citation_index}]" if s.citation_index else ""
            wrap(f"{s.index+1}. {s.text}{tag}")
            wrap(f"Similarity: {s.plag_score}% | Confidence: {s.confidence}")
            if s.snippet:
                wrap(f"Match: {s.snippet}", size=9, pad=12)
            wrap(f"Fix: {s.fix}", size=9, pad=12)

    if analysis.citations:
        line("Sources", bold=True, size=12, pad=16)
        for i, u in enumerate(analysis.citations, 1):
            wrap(f"[{i}] {u}", size=9, pad=12)

    c.showPage(); c.save(); buf.seek(0)

    return StreamingResponse(
        buf, media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename=\"{req.filename}\"'}
    )

class DocxIn(BaseModel):
    text: str
    filename: str = "detector_report.docx"
    source_texts: Optional[List[str]] = None
    source_urls: Optional[List[str]] = None
    plag_threshold: int = 85

@app.post("/report.docx")
async def report_docx(req: DocxIn):
    analysis = await core(AnalyzeIn(
        text=req.text, source_texts=req.source_texts, source_urls=req.source_urls, plag_threshold=req.plag_threshold
    ))
    doc = Document()
    doc.add_heading('AI Text Detection Report', 0)
    p = doc.add_paragraph("This report is informational. Use human judgment and cite sources where appropriate.")
    p.runs[0].font.size = Pt(10)

    doc.add_heading('Overall Results', level=1)
    doc.add_paragraph(f"AI-likeness: {analysis.summary['ai_score']}%")
    doc.add_paragraph(f"Plagiarism risk: {analysis.summary['plagiarism_score']}%")
    doc.add_paragraph(f"Human-likeness: {analysis.summary['human_score']}%")

    doc.add_heading('AI-flagged sentences', level=1)
    ai_items = [s for s in analysis.sentences if s.label == "ai"]
    if not ai_items:
        doc.add_paragraph("None detected ✅")
    else:
        for s in ai_items:
            doc.add_paragraph(f"{s.index+1}. {s.text}")
            doc.add_paragraph(f"Reason: {s.reason} | Confidence: {s.confidence}")
            doc.add_paragraph(f"Fix: {s.fix}")

    doc.add_heading('Plagiarism-risk sentences', level=1)
    pl_items = [s for s in analysis.sentences if s.label == "plag"]
    if not pl_items:
        doc.add_paragraph("None detected ✅")
    else:
        for s in pl_items:
            tag = f" [{s.citation_index}]" if s.citation_index else ""
            doc.add_paragraph(f"{s.index+1}. {s.text}{tag}")
            doc.add_paragraph(f"Similarity: {s.plag_score}% | Confidence: {s.confidence}")
            if s.snippet:
                doc.add_paragraph(f"Match: {s.snippet}")
            doc.add_paragraph(f"Fix: {s.fix}")

    if analysis.citations:
        doc.add_heading('Sources', level=1)
        for i, u in enumerate(analysis.citations, 1):
            doc.add_paragraph(f"[{i}] {u}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename=\"{req.filename}\"'}
    )
